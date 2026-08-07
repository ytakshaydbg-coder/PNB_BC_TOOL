"""
formatter.py - Word Document Generation Module
Project: PNB BC Agent Passbook Formatter Pro

This module is responsible for mapping extracted customer data to a pre-formatted 
Microsoft Word template. It strictly preserves existing formatting (Font, Size, 
Calibri, Alignment, Spacing) by performing in-place text replacement within the 
document's underlying run elements.
"""

import os
import logging
from typing import Dict, Any
from docx import Document

# Configure local logger
logger = logging.getLogger(__name__)

def replace_text_in_paragraph(paragraph: Any, replacements: Dict[str, str]) -> None:
    """
    Replaces placeholder text in a paragraph while preserving the run-level formatting.
    Handles both single-run placeholders and the common python-docx split-run issue.

    Args:
        paragraph: A python-docx Paragraph object.
        replacements (Dict[str, str]): Dictionary mapping placeholders to actual values.
    """
    # Step 1: Attempt standard run-level replacement (preserves exact intra-paragraph formatting)
    for run in paragraph.runs:
        for placeholder, value in replacements.items():
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(value))
                
    # Step 2: Fallback for split runs
    # Sometimes Word splits a placeholder like "{{NAME}}" across multiple runs 
    # (e.g., "{", "{NAME", "}}"). If the placeholder is still in the paragraph text,
    # we consolidate it into the first run to preserve the primary paragraph style.
    for placeholder, value in replacements.items():
        if placeholder in paragraph.text:
            # Reconstruct the full text with the replacement
            full_text = paragraph.text.replace(placeholder, str(value))
            if paragraph.runs:
                # Assign the new text to the first run to keep the baseline style
                paragraph.runs[0].text = full_text
                # Clear out the subsequent runs to avoid duplication
                for run in paragraph.runs[1:]:
                    run.text = ""

def generate_word_document(data: Dict[str, str], template_path: str = "template.docx", output_filename: str = "PNB_Passbook.docx") -> str:
    """
    Opens the specified Word template, replaces placeholders with customer data,
    and saves the generated file to the outputs directory.

    Args:
        data (Dict[str, str]): The parsed customer data dictionary.
        template_path (str): Path to the base Word template. Defaults to "template.docx".
        output_filename (str): Name of the output file. Defaults to "PNB_Passbook.docx".

    Returns:
        str: The file path of the successfully generated document, or an empty string on failure.
    """
    # Mapping the incoming dictionary keys to their exact document placeholders
    replacements: Dict[str, str] = {
        "{{NAME}}": data.get("NAME", ""),
        "{{CIF}}": data.get("CIF", ""),
        "{{ACCOUNT_NO}}": data.get("ACCOUNT_NO", ""),
        "{{AADHAAR}}": data.get("AADHAAR", ""),
        "{{MODE}}": data.get("MODE", ""),
        "{{OPEN_DATE}}": data.get("OPEN_DATE", ""),
        "{{ADDRESS}}": data.get("ADDRESS", ""),
        "{{PIN}}": data.get("PIN", ""),
        "{{NOMINATION}}": data.get("NOMINATION", ""),
        "{{ISSUE_DATE}}": data.get("ISSUE_DATE", "")
    }

    output_dir = "outputs"
    output_path = os.path.join(output_dir, output_filename)

    try:
        # Ensure output directory exists
        os.makedirs(output_dir, exist_ok=True)

        # Load the Word template
        if not os.path.exists(template_path):
            logger.error(f"Template file '{template_path}' not found.")
            raise FileNotFoundError(f"Template file '{template_path}' not found. Please ensure it exists in the root directory.")

        doc = Document(template_path)

        # 1. Replace placeholders in standard paragraphs
        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph, replacements)

        # 2. Replace placeholders inside tables (Passbook templates often use tables for structure)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_text_in_paragraph(paragraph, replacements)

        # Save the populated document
        doc.save(output_path)
        logger.info(f"Successfully generated formatted passbook: {output_path}")
        
        return output_path

    except Exception as e:
        logger.error(f"Failed to generate Word document: {str(e)}", exc_info=True)
        return ""